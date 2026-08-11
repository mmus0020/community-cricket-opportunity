from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import csv, math, os

OUT = Path('/mnt/data/phase18_final_report')
OUT.mkdir(exist_ok=True)
DOCX = OUT/'Community_Cricket_Opportunity_Final_Analytical_Report.docx'

# -------- helpers --------
def read_csv(path):
    with open(path, encoding='utf-8-sig') as f:
        return list(csv.DictReader(f))

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None:
            node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def add_caption(doc, text):
    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(90,90,90)
    return p

def add_key_box(doc, title, body, fill='E8F3F1'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # keep callout box on one page
    trPr = t.rows[0]._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)
    cell = t.cell(0,0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + '\n')
    r.bold = True; r.font.color.rgb = RGBColor(24,50,74); r.font.size=Pt(11)
    r2 = p.add_run(body); r2.font.size=Pt(10)
    return t

def add_bullet(doc, text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p

def add_small_source(doc, text):
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    return p

def add_image(doc, path, width=6.4, caption=None):
    path=Path(path)
    if path.exists():
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        if caption: add_caption(doc, caption)

def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table=doc.add_table(rows=1, cols=len(headers))
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.style='Table Grid'
    hdr=table.rows[0]
    set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        cell=hdr.cells[i]; shade_cell(cell,'18324A'); set_cell_margins(cell)
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(h)); r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(font_size)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]
            r=p.add_run(str(val)); r.font.size=Pt(font_size)
        if len(table.rows)%2==0:
            for c in cells: shade_cell(c,'F5F7F9')
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths):
                row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return table

# -------- load data --------
top20=read_csv('/mnt/data/phase12_opportunity_model/tables/opportunity_model_v1_top20.csv')
segments=read_csv('/mnt/data/phase13_rank_and_segment/tables/opportunity_model_v1_segmented.csv')
cases=read_csv('/mnt/data/phase14_top_opportunities/tables/phase14_case_study_profiles.csv')
nat_recs=read_csv('/mnt/data/phase15_strategic_recommendations/tables/national_recommendations.csv')
case_recs=read_csv('/mnt/data/phase15_strategic_recommendations/tables/case_specific_recommendations.csv')
limitations=read_csv('/mnt/data/phase16_validation_limitations/tables/limitations_register.csv')
validations=read_csv('/mnt/data/phase16_validation_limitations/tables/validation_register.csv')
state=read_csv('/mnt/data/phase10_eda/tables/state_eda_summary.csv')
corridors=read_csv('/mnt/data/phase11_geographic_analysis/tables/growth_corridor_summary.csv')
weights=read_csv('/mnt/data/phase12_opportunity_model/tables/weight_sensitivity_summary.csv')
source_local=read_csv('/mnt/data/phase14_top_opportunities/tables/phase14_source_register.csv')

# -------- document setup --------
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.65); sec.bottom_margin=Inches(0.65); sec.left_margin=Inches(0.75); sec.right_margin=Inches(0.75)

styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10)
styles['Normal'].paragraph_format.space_after=Pt(6)
for name,size,color in [('Title',28,'18324A'),('Heading 1',18,'18324A'),('Heading 2',13,'1F6F78'),('Heading 3',11,'18324A')]:
    s=styles[name]; s.font.name='Aptos Display' if name!='Normal' else 'Aptos'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(10); s.paragraph_format.space_after=Pt(5)
styles['Caption'].font.name='Aptos'; styles['Caption'].font.size=Pt(9); styles['Caption'].font.italic=True

# custom subtitle
if 'Report Subtitle' not in styles:
    st=styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name='Aptos'; st.font.size=Pt(13); st.font.color.rgb=RGBColor(80,95,110)

# header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text='Community Cricket Opportunity Analysis'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size=Pt(8); hp.runs[0].font.color.rgb=RGBColor(110,120,130)
    fp=section.footer.paragraphs[0]; fp.text='Public-data decision-support prototype | August 2026'; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size=Pt(8); fp.runs[0].font.color.rgb=RGBColor(120,120,120)

# -------- cover --------
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(60)
r=p.add_run('WHERE ARE AUSTRALIA\'S NEXT\nCOMMUNITY CRICKETERS?')
r.bold=True; r.font.size=Pt(27); r.font.color.rgb=RGBColor(24,50,74)
p=doc.add_paragraph(style='Report Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('A Data-Driven Analysis of Community Cricket Participation Growth Opportunities Across Australia')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28)
r=p.add_run('Final Analytical Report'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(31,111,120)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Prepared by Md Musa\nAugust 2026').font.size=Pt(11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(45)
r=p.add_run('PORTFOLIO / APPLICATION PROTOTYPE'); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(183,92,75)
add_key_box(doc,'Important interpretation boundary','This report identifies relative demographic Opportunity Potential for community cricket acquisition. It does not directly measure cricket demand, participation penetration, or Cricket System Readiness.', 'FFF4E5')
doc.add_page_break()

# -------- executive summary --------
doc.add_heading('Executive Summary', level=1)
p=doc.add_paragraph()
p.add_run('Purpose. ').bold=True
p.add_run('This project asks a practical participation-strategy question: which Australian communities warrant deeper investigation for future community cricket acquisition, and what evidence should determine the appropriate response?')
p=doc.add_paragraph()
p.add_run('Approach. ').bold=True
p.add_run('Using Australian Bureau of Statistics population, age/sex and socioeconomic data, Cricket Australia strategy and participation context, state cricket evidence and local case-study sources, the project builds a transparent LGA-level Community Cricket Opportunity Potential model. The model deliberately separates demographic opportunity from Cricket System Readiness.')

add_key_box(doc,'Headline finding','The strongest demographic growth signals are concentrated in a small number of metropolitan growth systems. The top 20 LGAs for absolute growth in the 5-14 population added 106,387 children between 2019 and 2024; 94.6% of that Top-20 increase was concentrated in four broad growth systems: outer Melbourne, western/south-west Sydney, South East Queensland and the Perth growth belt.')

doc.add_heading('Key findings', level=2)
for x in [
    'Market size, youth concentration and junior growth provide distinct information. Raw total population, junior population and girls\' junior population are too highly correlated to be weighted independently.',
    'The balanced Model V1 ranks Wyndham, Melton, Blacktown, Ipswich, Camden, Casey, Logan, Hume and The Hills among the strongest demographic opportunity areas nationally.',
    'The Top 20 is highly stable under alternative weighting scenarios and reasonable eligibility thresholds, indicating that the shortlist is not an artefact of one arbitrary specification.',
    'External case studies support different management responses: Blacktown shows a plausible capacity-gap pattern; Ipswich shows growth alongside new capacity; Armadale requires stronger participation validation before scaling.',
    'The recommended national decision process is: Opportunity screening -> Readiness diagnosis -> Problem classification -> Intervention -> Outcome measurement.'
]: add_bullet(doc,x)

add_image(doc,'/mnt/data/phase17_powerbi_dashboard/Dashboard_Executive_Overview_Preview.png',6.8,'Figure 1. Executive dashboard prototype developed from the Opportunity Potential model.')

# -------- contents --------
doc.add_page_break(); doc.add_heading('Contents',level=1)
contents=[
('1','Business Problem and Strategic Context'),('2','Research Question and Scope'),('3','Data and Methodology'),('4','Exploratory Findings'),('5','Geographic Analysis'),('6','Opportunity Potential Model'),('7','Ranking and Segmentation'),('8','Priority Case Studies'),('9','Strategic Recommendations'),('10','Validation and Limitations'),('11','Power BI Dashboard Design'),('12','Conclusion'),('Appendix A','Top 20 Model V1 Rankings'),('Appendix B','Model Specification'),('Appendix C','Validation and Limitations Register'),('Appendix D','Selected Sources')]
for n,t in contents:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    p.add_run(f'{n}. ').bold=True; p.add_run(t)

# -------- 1 --------
doc.add_page_break(); doc.add_heading('1. Business Problem and Strategic Context',level=1)
p=doc.add_paragraph('The Community Cricket Data & Insights Manager role requires more than reporting participation totals. The underlying business problem is how to integrate fragmented participation, customer, demographic, research and operational evidence so Australian Cricket can decide where to invest, whom to target, how to retain participants and how to measure whether participation strategies are working.')

doc.add_heading('Strategic context',level=2)
for x in [
    'Australian Cricket\'s 2022-2027 strategy places strong emphasis on participation growth, junior participation, women and girls, multicultural inclusion and a sport-for-all vision.',
    'Published participation ambitions include 210,000 registered participants aged 5-12 by 2027, including 60,000 girls, and more than 850,000 overall registered participants by 2027.',
    'The Women & Girls Action Plan 2024-2034 sets a longer-term ambition to grow girls aged 5-12 from roughly 25,000 to 100,000 by 2034.',
    'The Multicultural Action Plan focuses on creating opportunities, reducing barriers and increasing South Asian participation, reinforcing the need for place-based, community-specific decision support.'
]: add_bullet(doc,x)
add_small_source(doc,'Strategic source set: Cricket Australia strategy, Women & Girls Action Plan, Multicultural Action Plan, Australian Cricket Census and state association participation reporting.')

# -------- 2 --------
doc.add_heading('2. Research Question and Scope',level=1)
add_key_box(doc,'Primary research question','Which Australian communities show the strongest potential for community cricket participation growth, and what factors should Cricket Australia and State/Territory associations consider when prioritising participation-acquisition initiatives?')
p=doc.add_paragraph('The primary unit of analysis is the Local Government Area (LGA). The analysis covers all Australian states and territories, with 2025 LGA geography used as the master structure. The model focuses on general acquisition opportunity, with particular strategic attention to junior participation, women and girls, multicultural growth and local delivery capacity.')

doc.add_heading('What is intentionally out of scope',level=2)
for x in ['Individual churn prediction','Professional cricket performance','Exact ROI or participant forecasts','Causal claims about demographic characteristics','Direct claims of low participation penetration where LGA registration data are unavailable']:
    add_bullet(doc,x)

# -------- 3 --------
doc.add_heading('3. Data and Methodology',level=1)
doc.add_heading('3.1 Data sources',level=2)
rows=[
['ABS Regional Population','2025 / time series','Current population, one-year and five-year population growth'],
['ABS Regional Population by Age and Sex','2024 / 2001-2024','Current 5-14 population, girls 5-14, historical junior growth'],
['ABS SEIFA','2021','Socioeconomic context'],
['Cricket Australia / State Associations','Various','Strategy, participation benchmarks and contextual validation'],
['Council / local cricket sources','2024-2026','Case-study evidence on capacity, infrastructure and programs'],
['Power BI-ready model output','Phase 17','Dashboard implementation dataset and DAX measures']]
add_table(doc,['Source','Reference period','Use in analysis'],rows,widths=[2.0,1.3,3.1],font_size=8.5)

doc.add_heading('3.2 Geographic and age decisions',level=2)
p=doc.add_paragraph('The master dataset contains 548 current LGA records. Age data are available for 546, with East Arnhem and Groote Archipelago affected by a documented 2025 geographic restructure. For national Model V1 ranking, 365 LGAs meet the minimum-population guardrails and valid-age criteria.')
p=doc.add_paragraph('The latest public ABS age workbook provides five-year bands rather than exact single-year ages. The model therefore uses the directly observed 5-14 population as the current junior-market measure rather than constructing a false-precision 5-12 estimate.')

doc.add_heading('3.3 Analytical architecture',level=2)
add_key_box(doc,'Core design principle','Opportunity Potential and Cricket System Readiness remain separate. A community can have strong demographic opportunity but insufficient local capacity, or strong cricket supply but weaker future demographic growth.')
for x in [
    'Market Size - current scale of the 5-14 population.',
    'Youth Concentration - share of the local population aged 5-14.',
    'Growth Momentum - combination of absolute and percentage growth in the 5-14 population from 2019 to 2024.',
    'Context - SEIFA and geographic factors used for interpretation rather than direct scoring.',
    'System Readiness - intended future layer covering clubs, programs, facilities, capacity and workforce.'
]: add_bullet(doc,x)

# -------- 4 --------
doc.add_heading('4. Exploratory Findings',level=1)
p=doc.add_paragraph('Exploratory data analysis materially simplified the model. The raw absolute population measures are highly redundant: total population and 5-14 population have a Spearman rank correlation of approximately 0.993, while girls 5-14 and total 5-14 counts are effectively the same ranking signal. This means a model that weighted each of those counts independently would substantially double-count LGA size.')
add_image(doc,'/mnt/data/phase10_eda/charts/01_population_vs_junior_market.png',6.4,'Figure 2. LGA size strongly determines absolute junior-market size.')
p=doc.add_paragraph('Growth is different. Overall population growth and junior growth are related but not interchangeable (Spearman approximately 0.57), supporting a separate Growth Momentum pillar. Large current junior markets can be stagnant or shrinking, while smaller markets can be expanding rapidly.')
add_image(doc,'/mnt/data/phase10_eda/charts/02_population_growth_vs_junior_growth.png',6.4,'Figure 3. Total population growth and junior growth provide related but distinct signals.')

doc.add_heading('4.1 Examples of strong absolute junior growth',level=2)
top_abs=read_csv('/mnt/data/phase10_eda/tables/top20_junior_absolute_growth.csv')[:10]
rows=[[r['rank'],r['LGA_NAME'],r['STATE'],f"{float(r['AGE_5_14_CHANGE_5Y']):,.0f}",f"{float(r['AGE_5_14_GROWTH_5Y_PCT']):.1f}%"] for r in top_abs]
add_table(doc,['Rank','LGA','State','Additional 5-14','Growth %'],rows,widths=[0.5,1.7,1.3,1.4,1.0])

# -------- 5 --------
doc.add_heading('5. Geographic Analysis',level=1)
p=doc.add_paragraph('The strongest absolute increases in the junior market are geographically concentrated rather than scattered randomly. The national Top 20 for absolute 5-14 growth collectively added 106,387 children between 2019 and 2024. Four broad growth systems account for approximately 100,671 of that increase (94.6%).')
rows=[]
for r in corridors[:4]:
    rows.append([r['geographic_cluster'],r['top20_lga_count'],f"{float(r['additional_age_5_14_2019_24']):,.0f}",r['member_lgas']])
add_table(doc,['Growth system','Top-20 LGAs','Additional 5-14','Member LGAs'],rows,widths=[1.8,0.8,1.2,3.0],font_size=8)
add_image(doc,'/mnt/data/phase11_geographic_analysis/charts/01_growth_corridor_contribution.png',6.4,'Figure 4. Contribution of the four major growth systems to Top-20 junior-market growth.')
add_key_box(doc,'Management implication','Participation planning should not treat high-growth LGAs only as isolated councils. Adjacent high-growth areas can form shared participation systems, suggesting a role for corridor-level club, facility, school and program planning.')

# -------- 6 --------
doc.add_heading('6. Community Cricket Opportunity Potential Model',level=1)
p=doc.add_paragraph('Model V1 converts the Phase 10 evidence into a transparent national prioritisation index. Every input is transformed to a national percentile among eligible LGAs. Because there is no LGA-level outcome variable with which to estimate empirically optimal weights, the model uses equal weights across three conceptually distinct pillars and then tests sensitivity to alternative scenarios.')
rows=[
['Market Size','AGE_5_14_2024','33.33%','Current scale of the addressable junior market'],
['Youth Concentration','AGE_5_14_SHARE_TOTAL_PCT','33.33%','Young demographic structure, reducing large-LGA bias'],
['Growth Momentum','50% absolute growth percentile + 50% percentage growth percentile','33.33%','Balances growth scale and growth rate']]
add_table(doc,['Pillar','Input','Weight','Interpretation'],rows,widths=[1.3,2.5,0.8,2.5],font_size=8.5)
add_key_box(doc,'Model formula','Opportunity Potential = (Market Size percentile + Youth Concentration percentile + Growth Momentum percentile) / 3')
add_image(doc,'/mnt/data/phase12_opportunity_model/charts/01_model_v1_top20.png',6.5,'Figure 5. Model V1 Top 20 Community Cricket Opportunity Potential LGAs.')

# top 10 table
rows=[]
for r in top20[:10]:
    rows.append([r['RANK_BALANCED'],r['LGA_NAME'],r['STATE'],f"{float(r['SCORE_BALANCED']):.2f}",f"{float(r['AGE_5_14_2024']):,.0f}",f"{float(r['AGE_5_14_GROWTH_5Y_PCT']):.1f}%"])
add_table(doc,['Rank','LGA','State','Score','5-14 population','5-year growth'],rows,widths=[0.45,1.4,1.3,0.7,1.2,1.0])

doc.add_heading('6.1 Robustness',level=2)
p=doc.add_paragraph('The Top 20 is highly stable under alternative weights. The Growth-led model preserves all 20, the Scale-led model preserves 18, and the Intensity-led model preserves 19. Alternative eligibility thresholds also retain 19-20 of the baseline Top 20. This supports the interpretation that the leading candidate set is driven by the underlying data structure rather than a narrow modelling choice.')
add_image(doc,'/mnt/data/phase16_validation_limitations/charts/01_weight_sensitivity_top20.png',6.2,'Figure 6. Top-20 overlap under alternative weighting scenarios.')

# -------- 7 --------
doc.add_heading('7. Ranking and Strategic Segmentation',level=1)
p=doc.add_paragraph('A ranking answers where; segmentation begins to answer what kind of opportunity is present. The model therefore groups LGAs based on the pattern of Market Size, Youth Concentration and Growth Momentum rather than treating every high score as the same business problem.')
segments_desc=[
['Balanced High-Opportunity Market','Strong scale, youth intensity and growth','Priority acquisition investigation; test capacity before scaling'],
['Large Growth Market','Large junior market plus strong growth','Scale acquisition where delivery capacity exists'],
['Emerging Young Growth Market','Young and fast-growing, smaller current scale','Build delivery capacity early'],
['Young High-Intensity Market','Young demographic structure, less scale/growth','Validate demand and access'],
['Large Established / Slower-Growth','Large current market, weaker growth','Retention, conversion and utilisation'],
['Rapid Growth Market','Strong growth from a smaller base','Seed entry pathways and monitor'],
['Mixed / Moderate Opportunity','No dominant pillar profile','Selective local investigation']]
add_table(doc,['Segment','Typical profile','Implication'],segments_desc,widths=[2.0,2.3,2.5],font_size=8.2)
add_image(doc,'/mnt/data/phase13_rank_and_segment/charts/01_market_vs_growth_segments.png',6.5,'Figure 7. Strategic opportunity segments separate current scale from growth momentum.')

# -------- 8 --------
doc.add_heading('8. Priority Case Studies',level=1)
p=doc.add_paragraph('Five cases were selected to represent different growth systems and analytical situations rather than simply taking the national top five. The purpose was to test whether local cricket and infrastructure evidence supports, complicates or contradicts Model V1.')
for c in cases:
    doc.add_heading(f"8.{cases.index(c)+1} {c['LGA_NAME']}, {c['STATE']}",level=2)
    p=doc.add_paragraph()
    p.add_run(f"Model rank #{c['model_rank']} | Opportunity score {float(c['opportunity_score']):.2f} | ").bold=True
    p.add_run(f"5-14 population {int(float(c['age_5_14_2024'])):,} | 2019-24 change {int(float(c['age_5_14_change_2019_24'])):+,} ({float(c['age_5_14_growth_2019_24_pct']):.1f}%).")
    p=doc.add_paragraph(); p.add_run('Local evidence. ').bold=True; p.add_run(c['local_evidence'])
    p=doc.add_paragraph(); p.add_run('Working interpretation. ').bold=True; p.add_run(c['phase14_interpretation'])
    p=doc.add_paragraph(); p.add_run('Priority management question. ').bold=True; p.add_run(c['next_question'])
add_image(doc,'/mnt/data/phase14_top_opportunities/charts/01_case_study_pillar_profiles.png',6.3,'Figure 8. Pillar profiles for the five Phase 14 case studies.')

# -------- 9 --------
doc.add_heading('9. Strategic Recommendations',level=1)
add_image(doc,'/mnt/data/phase15_strategic_recommendations/charts/01_recommended_decision_process.png',6.4,'Figure 9. Recommended decision process from opportunity screening to outcome measurement.')
add_key_box(doc,'Core recommendation','Use demographic Opportunity Potential to decide where to investigate, cricket-system evidence to diagnose the local constraint, and intervention-specific KPIs to determine what action works.')

doc.add_heading('9.1 National recommendations',level=2)
for r in nat_recs:
    p=doc.add_paragraph()
    p.add_run(r['id']+' - '+r['recommendation']+'. ').bold=True
    p.add_run(r['action'])

doc.add_heading('9.2 Strategic response framework',level=2)
rows=[
['High opportunity + capacity pressure','Build Capacity First / Protect Access','Audit facilities/club capacity; quantify unmet demand; protect access before acquisition campaigns'],
['High opportunity + new/available capacity','Activate / Accelerate','Expand Blast/junior/girls pathways and school-to-club conversion'],
['High opportunity + weak demand evidence','Validate Before Scaling','Use registrations, club/school evidence and pilots before major investment'],
['Large market + slower growth','Retain / Convert','Prioritise participant return, churn and pathway conversion'],
['Smaller market + rapid growth','Seed and Monitor','Establish low-cost entry pathways and reserve future capacity']]
add_table(doc,['Evidence pattern','Response','Action'],rows,widths=[2.1,1.8,3.0],font_size=8.1)

# -------- 10 --------
doc.add_heading('10. Validation and Limitations',level=1)
p=doc.add_paragraph('The model is suitable as a transparent public-data prioritisation prototype, but its validation is strongest for data integrity, construct logic and robustness - not for prediction. There is no future LGA-level new-registration outcome against which Model V1 can be tested.')
rows=[]
for v in validations:
    rows.append([v['validation_dimension'],v['assessment'],v['interpretation']])
add_table(doc,['Validation dimension','Assessment','Interpretation'],rows,widths=[2.2,1.2,3.6],font_size=7.8)

doc.add_heading('10.1 Critical claim boundaries',level=2)
add_key_box(doc,'The model CAN support','Relative statements such as: an LGA combines a large junior market, a young population structure and strong recent junior growth compared with other eligible Australian LGAs.')
add_key_box(doc,'The model CANNOT support','Claims that an LGA has the highest untapped cricket demand, low participation penetration, a predicted number of future registrations, or a guaranteed return from a specific facility/program investment.','FFF4E5')

doc.add_heading('10.2 Highest-severity limitations',level=2)
crit=[x for x in limitations if x['severity'] in ('Critical','High')]
for x in crit:
    p=doc.add_paragraph(style='List Bullet')
    p.add_run(x['limitation']+': ').bold=True; p.add_run(x['issue'])
add_image(doc,'/mnt/data/phase16_validation_limitations/charts/03_validation_maturity.png',6.2,'Figure 10. Validation is strongest for robustness and weakest for outcome/predictive validity.')

# -------- 11 --------
doc.add_heading('11. Power BI Dashboard Design',level=1)
p=doc.add_paragraph('The final analytical workflow has been translated into a Power BI-ready dashboard package. The dashboard is designed for a hiring manager or participation leader to move from national screening to LGA-level diagnosis while keeping the model limitation visible.')
for x in [
    'Executive Overview - KPIs, Top 15 ranking, opportunity tiers and scale-vs-growth view.',
    'Geographic Opportunity - LGA choropleth using official ABS 2025 boundaries.',
    'LGA Deep Dive - rank, score, current junior market, growth and pillar profile.',
    'Strategic Segmentation - why each LGA ranks highly.',
    'Case Studies - Wyndham, Blacktown, Ipswich, Armadale and Maitland.',
    'Validation & Limitations - sensitivity tests, limitations and claim boundaries.'
]: add_bullet(doc,x)
add_image(doc,'/mnt/data/phase17_powerbi_dashboard/Dashboard_Executive_Overview_Preview.png',6.7,'Figure 11. Static prototype of the dashboard Executive Overview page.')

# -------- 12 --------
doc.add_heading('12. Conclusion',level=1)
p=doc.add_paragraph('The strongest contribution of this project is not the specific rank order of LGAs. It is the decision structure created around the ranking. Public demographic data can identify where participation conditions justify deeper investigation; local cricket evidence then determines whether the relevant problem is acquisition, capacity, retention, conversion or demand validation.')
add_key_box(doc,'Final conclusion','The defensible use of Model V1 is to identify where Australian Cricket should investigate first. A production-grade version should combine Opportunity Potential with PlayHQ registrations, club/program capacity, facilities, waitlists, volunteer/workforce data and participant research to build the full Opportunity x Readiness system.')
p=doc.add_paragraph('This prototype demonstrates a complete analytical workflow: strategic problem definition, public-data audit, data engineering, exploratory analysis, geographic interpretation, transparent modelling, sensitivity testing, case-study validation, strategic recommendations and dashboard implementation design.')

# -------- appendices --------
doc.add_page_break(); doc.add_heading('Appendix A. Top 20 Model V1 Rankings',level=1)
rows=[]
for r in top20:
    rows.append([r['RANK_BALANCED'],r['LGA_NAME'],r['STATE'],f"{float(r['SCORE_BALANCED']):.2f}",f"{float(r['SCORE_MARKET_SIZE']):.1f}",f"{float(r['SCORE_YOUTH_INTENSITY']):.1f}",f"{float(r['SCORE_GROWTH_MOMENTUM']):.1f}"])
add_table(doc,['Rank','LGA','State','Score','Market','Youth','Growth'],rows,widths=[0.45,1.4,1.2,0.7,0.8,0.8,0.8],font_size=7.8)

doc.add_heading('Appendix B. Model Specification',level=1)
for x in [
    'Eligibility: valid age data, current total population >= 5,000, 2019 5-14 population >= 500, residual unincorporated records excluded except ACT.',
    'Normalisation: average-rank percentile scores among eligible LGAs.',
    'Market Size = percentile of AGE_5_14_2024.',
    'Youth Concentration = percentile of AGE_5_14_SHARE_TOTAL_PCT.',
    'Growth Momentum = average of percentiles for AGE_5_14_CHANGE_5Y and AGE_5_14_GROWTH_5Y_PCT.',
    'Balanced Opportunity Potential = equal average of Market Size, Youth Concentration and Growth Momentum.',
    'SEIFA is contextual only. Cultural strategic alignment and System Readiness are excluded from Model V1 because those layers were not fully integrated.'
]: add_bullet(doc,x)

doc.add_heading('Appendix C. Validation and Limitations Register',level=1)
rows=[]
for x in limitations:
    rows.append([x['severity'],x['limitation'],x['risk_to_interpretation'],x['mitigation']])
add_table(doc,['Severity','Limitation','Risk','Mitigation'],rows,widths=[0.8,1.6,2.6,2.3],font_size=7.2)

doc.add_heading('Appendix D. Selected Sources',level=1)
refs=[
('Cricket Australia','Australian Cricket Strategy 2022-2027 / Governing the Game','https://www.cricket.com.au/governing-the-game/ca-strategy'),
('Cricket Australia','Women and Girls Action Plan 2024-2034','https://www.cricket.com.au/social-impact-and-sustainability/women-and-girls'),
('Cricket Australia','Multicultural Action Plan / Diversity and Inclusion','https://www.cricket.com.au/social-impact-and-sustainability/diversity-and-inclusion/multicultural'),
('Australian Bureau of Statistics','Regional Population, 2024-25','https://www.abs.gov.au/statistics/people/population/regional-population/latest-release'),
('Australian Bureau of Statistics','Regional Population by Age and Sex','https://www.abs.gov.au/statistics/people/population/regional-population-age-and-sex/latest-release'),
('Australian Bureau of Statistics','Socio-Economic Indexes for Areas (SEIFA), 2021','https://www.abs.gov.au/statistics/people/people-and-communities/socio-economic-indexes-areas-seifa-australia/latest-release'),
('Australian Sports Commission','AusPlay results and data tables','https://www.ausport.gov.au/clearinghouse/research/ausplay/results'),
('Cricket Australia','2022-23 Australian Cricket Census','https://www.cricket.com.au/news/3646243/cricket-census-reveals-cricket-participation-continues-to-grow'),
]
for r in source_local:
    refs.append((r['source_org'],r['evidence'],r['url']))
for org,title,url in refs:
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(3)
    p.add_run(org+'. ').bold=True; p.add_run(title+'. '); rr=p.add_run(url); rr.font.size=Pt(8.5); rr.font.color.rgb=RGBColor(31,111,120)

# final note
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('End of report'); r.italic=True; r.font.color.rgb=RGBColor(120,120,120)

# save
# set core properties
doc.core_properties.title="Where Are Australia's Next Community Cricketers?"
doc.core_properties.subject='Community cricket participation opportunity analysis'
doc.core_properties.author='Md Musa'
doc.core_properties.keywords='cricket, participation, data analytics, Australia, community cricket, Power BI'
doc.save(DOCX)
print(DOCX)
